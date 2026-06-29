"""Word via python-docx (structured editing): read a document's paragraph/heading structure and
apply targeted edits (headings, paragraphs, edit/replace/delete text, images, tables, page breaks).
Saves the file, then opens it in Word as a live preview. Creates the document if it doesn't exist.
"""
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from axon.util import _result
from axon.office._base import _resolve_path
from axon.office.preview import ensure_closed, open_doc, save_with_retry


def _para(doc, i):
    paras = doc.paragraphs
    i = int(i)
    if i < 0 or i >= len(paras):
        raise IndexError(f"paragraph {i} out of range (document has {len(paras)} paragraphs)")
    return paras[i]


def _set_text(p, text):
    """Replace a paragraph's text while keeping its style."""
    for r in list(p.runs):
        r.text = ""
    (p.runs[0] if p.runs else p.add_run()).text = str(text)


def _style_runs(p, op):
    for r in p.runs:
        if op.get("bold") is not None:
            r.bold = bool(op["bold"])
        if op.get("italic") is not None:
            r.italic = bool(op["italic"])
        if op.get("size"):
            r.font.size = Pt(float(op["size"]))
        if op.get("color"):
            r.font.color.rgb = RGBColor.from_string(str(op["color"]).lstrip("#"))


def _replace_text(doc, find, repl):
    n = 0
    scopes = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                scopes += cell.paragraphs
    for p in scopes:
        if find in p.text:
            _set_text(p, p.text.replace(find, repl))
            n += 1
    return n


def _apply_op(doc, op):
    kind = (op.get("op") or "").lower()
    if kind == "add_heading":
        p = doc.add_heading(str(op.get("text", "")), level=int(op.get("level", 1)))
        return None
    if kind == "add_paragraph":
        p = doc.add_paragraph(str(op.get("text", "")), style=op.get("style") or None)
        _style_runs(p, op)
        return None
    if kind == "set_paragraph":
        p = _para(doc, op["index"])
        _set_text(p, op.get("text", ""))
        _style_runs(p, op)
        return None
    if kind == "delete_paragraph":
        p = _para(doc, op["index"])
        p._element.getparent().remove(p._element)
        return None
    if kind == "replace_text":
        n = _replace_text(doc, str(op.get("find", "")), str(op.get("replace", "")))
        return None if n else f"text not found: {op.get('find')!r}"
    if kind == "add_image":
        path = op.get("path")
        if not path or not os.path.exists(path):
            return f"image not found: {path}"
        doc.add_picture(path, width=Inches(float(op["width"])) if op.get("width") else None)
        return None
    if kind == "add_table":
        rows = op.get("rows") or []
        if not rows:
            return "table needs rows"
        ncols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=ncols)
        t.style = op.get("style", "Table Grid")
        for r, row in enumerate(rows):
            for c in range(ncols):
                t.cell(r, c).text = str(row[c]) if c < len(row) else ""
        return None
    if kind == "add_page_break":
        doc.add_page_break()
        return None
    return f"unknown op: {kind}"


def word_read(args):
    """Outline a document: each non-empty paragraph as [index] (style) text, plus table info."""
    path = args.get("path")
    if not path or not os.path.exists(path):
        return _result(f"No such file: {path}", True)
    try:
        doc = Document(path)
    except Exception as e:
        return _result(f"Could not open document: {e}", True)
    lines = [f"{os.path.basename(path)} — {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables"]
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        st = p.style.name if p.style else "Normal"
        tag = f"({st}) " if st and st != "Normal" else ""
        lines.append(f"[{i}] {tag}{p.text}")
    for ti, t in enumerate(doc.tables):
        lines.append(f"\n[table {ti}] {len(t.rows)}x{len(t.columns)}")
    return _result("\n".join(lines))


def word_edit(args):
    """Apply edit ops to a Word document (creating it if missing), save, and open it in Word.

    ops example:
      [{"op":"add_heading","text":"Proposal","level":0},
       {"op":"add_paragraph","text":"Dear team,"},
       {"op":"replace_text","find":"TBD","replace":"2026-07-01"},
       {"op":"add_table","rows":[["Item","Cost"],["Setup","$0"]]}]
    """
    path = _resolve_path(args.get("path"), "document", ".docx")
    ops = args.get("ops") or []
    if not ops:
        return _result("Provide ops: a list of edit operations (add_heading / add_paragraph / replace_text ...).", True)
    ensure_closed(path)
    try:
        doc = Document(path) if os.path.exists(path) else Document()
    except Exception as e:
        return _result(f"Could not open document: {e}", True)
    done, errors = 0, []
    for n, op in enumerate(ops, 1):
        try:
            err = _apply_op(doc, op)
            if err:
                errors.append(f"op {n} ({op.get('op')}): {err}")
            else:
                done += 1
        except Exception as e:
            errors.append(f"op {n} ({op.get('op')}): {e}")
    final, serr = save_with_retry(path, lambda p: doc.save(p))
    if serr is not None:
        return _result(f"Could not save document: {serr}", True)
    open_doc(final)
    msg = f"Applied {done}/{len(ops)} edits to {final}."
    if final != path:
        msg += " (Original was open/locked, so I saved an updated copy.)"
    if errors:
        msg += " Issues:\n" + "\n".join(errors)
    return _result(msg, bool(errors) and done == 0)
